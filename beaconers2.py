from flask import send_file
import pandas as pd
from flask import redirect, url_for
# from fuzzywuzzy import fuzz
# from fuzzywuzzy import process
from rapidfuzz import fuzz
from rapidfuzz import process
import os
import pytesseract
from pdf2image import convert_from_path
from flask import Flask, request, jsonify, send_from_directory, send_file, Response, url_for
from flask_cors import CORS, cross_origin

from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer

from datetime import datetime
from textblob import TextBlob
import json
import re
import PyPDF2
from docx import Document
from config import *
import subprocess

def generate_pdf(doc_path, path):
    print("inside PDF func")
    print(doc_path)
    print(path)
    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path

app = Flask(__name__)
app.config['CORS_HEADERS'] = 'Content-Type'
CORS(app, expose_headers=["Content-Disposition"])
# remove the "/" in the end of path
proot =BASE_DIR

if not os.path.exists(f'{proot}/temp'):
    with open(f'{proot}/temp', "w"): pass

name_dict = {'galaxy Siglight-AN.xlsx': 'Galaxy',
            'Jupiter Siglight-AN.xlsx': "Jupiter",
            'Mars- SIG Lite-AN.xlsx': "Mars",
            "Orit Siglight-AN.xlsx": "Orbit",
            "SIG Lite- Vendor Alpha-Anita.xlsx": "Alpha",
            "SIG Lite- Vendor Beta-Anita.xlsx": "Beta",
            "Vendor Sigma SIG Lite-Anita.xlsx": "Sigma",
            "Vendor Froce_SIG Lite _ AN.xlsx": "Force",
            "Solar-Siglight_AN.xlsx": "Solar"
            }


def merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j):
    
    new_data = {'Question': [], 'Answer': [], 'Risk': []}
    if ('Sheet' not in new_questions.columns) and ('Serial number' not in new_questions.columns):
        
        #print(new_questions["Question"][4])
        for i in range(len(column_13_data)):
            answer = None
            #print("question length and row length")
            #print(len(column_12_data))
            #print(len(column_13_data))
            value_12=column_12_data[i]
            value_13 = column_13_data[i]
            
            new_data['Question'].append(value_13)
            match_index = new_questions.index[((new_questions.index)+1) == value_12].tolist()
            #print("match index")
            #print(match_index)
            if match_index:
                
                answer = new_questions.loc[match_index[0], 'Answer']
                new_data['Answer'].append(answer)
                
            else:
                new_data['Answer'].append('Not Available')
            

            answer_str = str(answer)  

            if answer_str.lower() in ['yes', 'yes ', 'Yes', 'Yes ', "YES"]:
                risk = 'Green'
            elif answer_str.lower() in ['no', 'no ', 'No', 'No ', "NO"]:
                risk = 'Red'
            else:
                risk = ''

            new_data['Risk'].append(risk)
        #print(new_questions1)
        new_questions1 = pd.DataFrame(new_data)
        print("final result")
        print(new_questions1)
        return new_questions1
    else:    
        
        for i in range(len(column_13_data)):
            value_10 = column_10_data[i]
            value_11 = column_11_data[i]
            value_12 = column_12_data[i]
            value_13 = column_13_data[i]
            if value_10 not in ['yes', 'Yes', 'no', 'No', 'yes ', 'Yes ', 'no ', 'No ']:
                new_data['Question'].append(value_13)
                new_data['Answer'].append('Not Available')
                new_data['Risk'].append('')
                continue
            if j==10:
                match = new_questions[(new_questions['Serial number'] == str(value_12))]
            elif j==6:
                match = new_questions[(new_questions['Sheet'] == value_11) & ((new_questions['Serial number']) == value_12-3)]   
            else:
                match = new_questions[(new_questions['Sheet'] == value_11) & (new_questions['Serial number'] == value_12)]
            if not match.empty:
                answer = str(match.iloc[0]['Answer'])
        
                if answer.lower() in ['yes', 'yes ', 'Yes', 'Yes ', 'YES']:
                    risk = 'Green'
           
                elif answer.lower() in ['no', 'no ', 'No', 'No ', 'NO']:
                    risk = 'Red'
                else:
                    risk = ''
                new_data['Question'].append(value_13)
                new_data['Answer'].append(answer)
                new_data['Risk'].append(risk)
        new_questions1 = pd.DataFrame(new_data)
        print("final result")
        print(new_questions1)
        return new_questions1


def assessment_details(sheets, file_path,filename,basefile):
    sheets.remove("Assessment details")
    dfs = pd.read_excel(file_path, sheet_name=sheets, usecols=["Serial number", "Question", "Actual answers"])
    new_questions = pd.DataFrame(columns=["Serial number", "Question", "Answer", "Sheet"])

    for sheet, df in dfs.items():
        sheet_name = re.search(r"^[A-Z]", sheet).group()
        df['Sheet'] = sheet_name

        df.rename(columns={"Actual answers": "Answer"}, inplace=True)
        df['Question'] = df['Question'].apply(lambda row: re.sub("^\d+\.\d+|^[A-Z]{2}\.\d+\.", "", row, count=1))

        new_questions = pd.concat([new_questions, df[["Serial number", "Question", "Answer", "Sheet"]]], ignore_index=True)
    if filename=="Vandor Sigma SIG Lite-Anita.xlsx":
        column_10_data = basefile.iloc[1:, 8].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 9].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 10].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=1)
    elif filename=="SIG Lite- Vandor Alpha-Anita.xlsx":
        column_10_data = basefile.iloc[1:, 11].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 12].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 13].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=2)
    elif filename=="SIG Lite- Vendor Beta-Anita.xlsx":
        column_10_data = basefile.iloc[1:, 14].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 15].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 16].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=3)
    elif filename=="Vendor Froce_SIG Lite _ AN.xlsx":
        column_10_data = basefile.iloc[1:, 17].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 18].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 19].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=4)
    elif filename=="galaxy Siglight-AN.xlsx":
        column_10_data = basefile.iloc[1:, 26].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 27].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 28].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=8)
    elif filename=="Mars- SIG Lite-AN.xlsx":
        column_10_data = basefile.iloc[1:, 29].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 30].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 31].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=9)
    return new_questions1

def cover_page_no_lite(sheets, file_path,filename, basefile):
    print("in cover page no lite")
    remove_sheet = ['Cover Page', 'Business Information', 'Documentation', 'Lite', 'Glossary', 'Version History', 'Formula Notes', 'Full', 'Copyright', 'Terms of Use']
    for sheet in remove_sheet:
        sheets.remove(sheet)
    df_list = []
    for sheet in sheets:
        df = pd.read_excel(file_path, sheet_name=sheet, index_col=None, usecols="C,D")
        df = df.reset_index()
        sheet_name = re.search(r"^[A-Z]", sheet).group()  # Extract the sheet name using regex
        df['Sheet'] = sheet_name
        df_list.append(df)
    
    new_questions_list = []
    for df in df_list:
        sheet_questions = pd.DataFrame({
            "Question": df['Unnamed: 2'][2:].to_list(),
            "Answer": df['Unnamed: 3'][2:].to_list(),
            "Sheet": df['Sheet'][2:].to_list()
        })
        sheet_questions['Serial number'] = sheet_questions.index + 1
        new_questions_list.append(sheet_questions)
    
    new_questions = pd.concat(new_questions_list)
    #new_questions.to_excel("hello.xlsx", index=False)
    #orbit file 
    column_10_data = basefile.iloc[1:, 23].reset_index(drop=True)
    column_11_data = basefile.iloc[1:, 24].reset_index(drop=True)
    column_12_data = basefile.iloc[1:, 25].reset_index(drop=True)
    column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
    
    new_questions1 = merge_sheet_answer(column_10_data, column_11_data, column_12_data, column_13_data, new_questions, j=6)
    
    
    return new_questions1


def cover_page(sheets, file_path,filename,basefile):
    df = pd.read_excel(file_path, sheet_name='Lite')
    for i, j in enumerate(df['Unnamed: 3']):
        if j == 'Response':
            if pd.isna(df['Unnamed: 3'][i+1]) and pd.isna(df['Unnamed: 3'][i+2]) and pd.isna(df['Unnamed: 3'][i+3]) and pd.isna(df['Unnamed: 3'][i+4]):
                return cover_page_no_lite(sheets, file_path,filename,basefile)
            else:
                new_questions = pd.DataFrame()
                new_questions['Question'] = df['Unnamed: 2'][2:]
                new_questions['Answer'] = df['Unnamed: 3'][2:]
                df['SIG Lite'] = df['SIG Lite'].str.replace('SL.', '', regex=True).str.strip()
                new_questions['Serial number'] = df['SIG Lite'][2:]
    
    #if filename=="Jupiter SIglight-AN.xlsx":
    column_10_data = basefile.iloc[1:, 32].reset_index(drop=True)
    column_11_data = basefile.iloc[1:, 33].reset_index(drop=True)
    column_12_data = basefile.iloc[1:, 34].reset_index(drop=True)
    column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
    new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=10)
    print("cover page")
    return new_questions1


def sheets_zero(file_path,filename,basefile):
    start_row = 0
    start_column = 0

    new_questions = pd.read_excel(file_path, header=None)

    for index, row in new_questions.iterrows():
        if 'Question' in str(row):
            start_row = index
            start_column = row[row.astype(str).str.contains('Question', case=False, na=False)].index.tolist()[0]
            # print(start_row)
            # print(start_column)
            break
        elif 'Questionnaire'in str(row):
            continue
    questions = new_questions.iloc[start_row+1:, start_column]
    questions = questions.str.replace(r"^\d+\.\d+|^[A-Z]{2}\.\d+\.", '', regex=True)
    #print(questions)

    for index, row in new_questions.iterrows():
        if 'Response' in str(row):
            start_row = index
            start_column = row[row.astype(str).str.contains('Response', case=False, na=False)].index.tolist()[0]
            break

    answer = new_questions.iloc[start_row+1:, start_column]

    # new_questions = new_questions[['Unnamed: 0', 'Unnamed: 2']][7:]
    new_questions = pd.DataFrame({"Question": questions, "Answer": answer})
    new_questions['Question'] = new_questions['Question'].apply(lambda row: re.sub("^\d+\.\d+", "", row, count=1))
    if filename=="galaxy Siglight-AN.xlsx":
        column_10_data = basefile.iloc[1:, 26].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 27].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 28].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=7)
    elif filename=="Solar-Siglight_AN.xlsx":
        column_10_data = basefile.iloc[1:, 20].reset_index(drop=True)
        column_11_data = basefile.iloc[1:, 21].reset_index(drop=True)
        column_12_data = basefile.iloc[1:, 22].reset_index(drop=True)
        column_13_data = basefile.iloc[1:, 2].reset_index(drop=True)
        new_questions1=merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j=5)
    return new_questions1


def perform_sentiment_analysis(data):
    def sentiment_analysis(answer):
        try:
            if answer.startswith("not") or answer.startswith("Not") or answer.startswith("NOT") or answer.startswith("not ") or answer.startswith("Not ") or answer.startswith("NOT "):
                sentiment_label = 'Not Available'
                polarity = None
                #subjectivity = None
            elif answer.startswith("Yes") or answer.startswith("yes") or answer.startswith("YES") or answer.startswith("Yes ") or answer.startswith("yes ") or answer.startswith("YES "):
                sentiment_label = 'Yes'
                polarity = None
                #subjectivity = None
            
            elif answer.startswith("No") or answer.startswith("NO") or answer.startswith("no") or answer.startswith("No ") or answer.startswith("NO ") or answer.startswith("no "):
                sentiment_label = 'No'
                polarity = None
                #subjectivity = None
            elif answer.startswith("N/A") or answer.startswith("n/a") or answer.startswith("Not Answered") or answer.startswith("NA") or answer.startswith("NOT ANSWERED") or answer.startswith("not answered") or answer.startswith("Not answered") or answer.startswith("na") or answer.startswith("N A"):
                sentiment_label='N/A'
                polarity=None
            else:
                blob = TextBlob(answer)
                polarity = blob.sentiment.polarity
                #subjectivity = blob.sentiment.subjectivity

                if polarity >= 0:
                    sentiment_label = 'Yes'
                elif polarity < 0:
                    sentiment_label = 'No'
        except:
            sentiment_label = 'Not Available'
            polarity = None

        return sentiment_label

    data[['Answer']] = data['Answer'].apply(sentiment_analysis).apply(pd.Series)

    return data


@app.route("/compare_questionnaires", methods=["GET", "POST"])
@cross_origin()
def compare_questionnaires():
    print("in this function")
    data = request.files['SIGlight']
    data.save(os.path.join(f'{proot}/temp/', data.filename))
    file_path = os.path.join(f'{proot}/temp/', data.filename)
    # fileobj = data.read()

    # Set up base template
    base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions = pd.read_excel(base_file)
    base_questions1 = pd.read_excel(base_file)
    base_questions= base_questions.iloc[1:, 2].reset_index(drop=True)
    #base_questions = base_questions['Unnamed: 4'][2:]
    base_questions = base_questions.apply(lambda row: re.sub("^\d+\.\d+", "", row, count=1))

    # Set up file to compare
    # new_file = file_path
    try:
        new_questions = pd.read_excel(file_path)
    except:
        return "Error in reading Excel file"
    file = pd.ExcelFile(file_path)
    sheets = file.sheet_names

    if len(sheets) > 1:
        if "Assessment details" in sheets:
            new_questions = assessment_details(sheets, file_path,data.filename,base_questions1)
        elif "Cover Page" in sheets:
            new_questions = cover_page(sheets, file_path,data.filename,base_questions1)
    else:
        new_questions = sheets_zero(file_path,data.filename,base_questions1)

    #print("HERE....")
    #print(new_questions)

    dict = {'Question': [],
            'matched_question': [],
            'Answer': []}
    # test_question_list = new_questions['Question'].to_list()
    # for row in base_questions:
    #     temp = process.extractOne(row, test_question_list)
    #     if temp[1] > 75:
    #         dict['Question'].append(row)
    #         if type(new_questions['Answer'][temp[2]]) is not float:
    #             dict['Answer'].append(new_questions["Answer"][temp[2]])
    #             dict['input_question'].append(test_question_list[temp[2]])
    #             test_question_list.remove(test_question_list[temp[2]])
    #         else:
    #             dict['Answer'].append("Not Answered")
    #             dict['input_question'].append("No match")
    #     else:
    #         dict['Question'].append(row)
    #         dict['input_question'].append(test_question_list[temp[2]])
    #         dict['Answer'].append("No")

    for row in base_questions:
        temp = process.extractOne(row, new_questions['Question'])
        if temp[1] > 85:
            dict['Question'].append(row)
            dict['matched_question'].append(new_questions['Question'])
            if type(new_questions['Answer'][temp[2]]) is not float:
                dict['Answer'].append(new_questions["Answer"][temp[2]])
            else:
                dict['Answer'].append("Not Answered")
        else:
            dict['Question'].append(row)
            dict['matched_question'].append(new_questions['Question'])
            dict['Answer'].append("No")
    print("function end")
    print(len(dict['Question']))
    print(len(dict['Answer']))

    df = pd.DataFrame(dict) 
    updated_df = perform_sentiment_analysis(df)
    print(type(updated_df['Question']))

    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # updated_df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()
    # # updated_df.to_excel(f"{proot}/output/output.xlsx")
    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))

    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    print([i.cells[1].text])
                    print(updated_df.index[updated_df['Question'] == i.cells[1].text])
                    # if not updated_df.index[updated_df['Question'] == i.cells[1].text].empty:
                    count = updated_df.index[updated_df['Question'] == i.cells[1].text][0]
                    c += 1
                    i.cells[2].text = updated_df['Answer'][count]

    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if 'Vendor Name' in run.text:
                run.text = run.text.replace(old_word, new_word)
                break
    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\")

    return json.dumps({'Question': updated_df['Question'].to_list(),
                       'Answer': updated_df['Answer'].to_list()}, indent=4)


def convert_pdf_to_text(pdf_file):
    # pdf_file = open("soc_2/Venus-SOC2-Type2 AN.pdf", "rb")
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    num_pages = len(pdf_reader.pages)
    txt_file = open("temp_pdf_text.txt", "w")
    for i in range(num_pages):
        page_text = pdf_reader.pages[i].extract_text()
        txt_file.write(page_text)
        # print(page_text)
    # pdf_file.close()
    txt_file.close()
    return "temp_pdf_text.txt"

stopwords= ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've",\
            "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', \
            'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their',\
            'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', \
            'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', \
            'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', \
            'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after',\
            'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further',\
            'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more',\
            'most', 'other', 'some', 'such', 'only', 'own', 'same', 'so', 'than', 'too', 'very', \
            's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', \
            've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn',\
            "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn',\
            "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", \
            'won', "won't", 'wouldn', "wouldn't"]

def clean_pdf(pdf_txt_path):
    with open(pdf_txt_path, 'r', encoding='utf-8') as file:
        texts = file.readlines()

    cleaned_lines = []
    for text in texts:
        text = text.strip()

        if not text:
            continue

        text = re.sub(r'\b\w\b', '', text)
        text = re.sub(r"\\", " ", text)
        text = re.sub(r"\/", " ", text)
        text = re.sub(r"[\n\t\-]", " ", text)
        text = re.sub(r"won't", "will not", text)
        text = re.sub(r"won’t", "will not", text)
        text = re.sub(r"can\'t", "can not", text)
        text = re.sub(r"can\’t", "can not", text)
        text = re.sub(r"n\'t", " not", text)
        text = re.sub(r"n\’t", " not", text)
        text = re.sub(r"\'re", " are", text)
        text = re.sub(r"\’re", " are", text)
        text = re.sub(r"\'s", " is", text)
        text = re.sub(r"\’s", " is", text)
        text = re.sub(r"\'d", " would", text)
        text = re.sub(r"\’d", " would", text)
        text = re.sub(r"\'ll", " will", text)
        text = re.sub(r"\’ll", " will", text)
        text = re.sub(r"\'t", " not", text)
        text = re.sub(r"\’t", " not", text)
        text = re.sub(r"\'ve", " have", text)
        text = re.sub(r"\’ve", " have", text)
        text = re.sub(r"\’m", " am", text)
        text = re.sub(r"\'m", " am", text)
        text = re.sub(r'[^A-Za-z_]', " ", text)
        text = re.sub(r"\s+", " ", text)
        text = ' '.join(e for e in text.split() if e.lower() not in stopwords)
        cleaned_lines.append(text.lower())

    cleaned_text = "\n".join(cleaned_lines)
    if os.path.exists(pdf_txt_path):
        os.remove(pdf_txt_path)
    return cleaned_text



@app.route("/compare_pdfs", methods=['POST', 'GET'])
@cross_origin()
def compare_pdfs():
    # Load the BERT model
    data = request.files['SOC2']
    # print(data.read())
    file_name = data.filename
    pdf_path = os.path.join(f"{proot}/temp/", file_name)
    # data.save(pdf_path)
    # output_text_file = os.path.join("output/", str(datetime.now()))

    pdf_txt_path = convert_pdf_to_text(data)
    text = clean_pdf(pdf_txt_path)
    # images = convert_from_path(pdf_path)
    # text = ""
    # for image in images:
        # text += pytesseract.image_to_string(image, lang='eng')
    # with open(output_text_file, 'w') as file:
        # file.write(text)

    model = SentenceTransformer('bert-base-nli-mean-tokens')

    # Load the Excel sheet for sheet1
    sheet1 = pd.read_excel(f'{proot}/base_files/template.xlsx', header=None)

    # with open('/content/Atlassian-Platform-SOC2-Type-2_30-Sep-2021-1.txt', 'r') as file:
    #     texts_sheet2 = file.readlines()
    #texts_sheet2 = [preprocess(line) for line in texts_sh2]
    # texts_sheet2 = preprocess(texts_sheet2)

  # text_sheet1: base template
    texts_sheet1 = sheet1.values.flatten().tolist()
    # text_sheet2: soc2 text
    texts_sheet2 = text.split('\n')


    # embeddings_sheet1 = model.encode(texts_sheet1)
    # embeddings_sheet2 = model.encode(texts_sheet2)

    matching_questions = []
    unmatched_questions = []

    for row in texts_sheet1:
        temp = process.extractOne(row, texts_sheet2)
        if temp[1] > 80:
            matching_questions.append((row, temp[1], temp[1], "Yes"))
        else:
            matching_questions.append((row, temp[1], temp[1], 'No'))

    # for i, embedding_sheet1 in enumerate(embeddings_sheet1):
    #     max_similarity = 0.0
    #     max_similarity_index = -1
    #     for j, embedding_sheet2 in enumerate(embeddings_sheet2):
    #         similarity = cosine_similarity([embedding_sheet1], [embedding_sheet2])[0][0]
    #         if similarity > max_similarity:
    #             max_similarity = similarity
    #             max_similarity_index = j

    #     if max_similarity_index != -1:
    #         if max_similarity > 0.75:
    #             matching_questions.append((sheet1.iloc[i, 0], texts_sheet2[max_similarity_index], max_similarity, 'Yes'))
    #         else:
    #             matching_questions.append((sheet1.iloc[i, 0], texts_sheet2[max_similarity_index], max_similarity, 'No'))

            # matching_questions.append((sheet1.iloc[i, 0], texts_sheet2[max_similarity_index], max_similarity, 'Yes'))
        # else:
        #     unmatched_questions.append(sheet1.iloc[i, 0])

    matching_df = pd.DataFrame(matching_questions, columns=['Question 1', 'Question 2', 'Similarity', 'Matching'])


    # unmatched_df = pd.DataFrame({'Question 1': unmatched_questions, 'Question 2': '', 'Similarity': '', 'Matching': 'No'})
    # updated_df = matching_df.append(unmatched_df, ignore_index=True)
    updated_df = {'Question': matching_df['Question 1'].to_list(),
                  'Answer': matching_df['Matching'].to_list()}
                  # 'Similarity': matching_df['Similarity'].to_list()}

    df = pd.DataFrame(updated_df)
    # excel_file = file_name.split(".")[0] + ".xlsx"
    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()
    print(df)
    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))

    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    print([i.cells[1].text])
                    print(df.index[df['Question'] == i.cells[1].text])
                    if not df.index[df['Question'] == i.cells[1].text].empty:
                        count = df.index[df['Question'] == i.cells[1].text][0]
                        c += 1
                        i.cells[2].text = df['Answer'][count]
                    else:
                        i.cells[2].text = "Not Answered"
                        c +=1
    
    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if 'Vendor Name' in run.text:
                run.text = run.text.replace(old_word, new_word)
                break


    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\")


    # matching_df.to_excel(os.path.join("output/", "output.xlsx"), index=False)
    # updated_df.to_excel(os.path.join("temp/", "output.xlsx"), index=False)
    return json.dumps(updated_df, indent=4)
    # print("Unmatched Questions:")
    # for question in unmatched_questions:
    #     print(question)


def update_excel_with_similarity(sheet1):
    data = request.files['SOC2']
    data2 = request.files['SIGlight']
    file_name = data.filename
    pdf_path = os.path.join(f"{proot}/temp/", file_name)
    data.save(pdf_path)
    output_text_file = os.path.join("output/", str(datetime.now()))
    #output_text_file = os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", str(datetime.now()))
    pdf_txt_path = convert_pdf_to_text(data)
    texts_sheet2 = clean_pdf(pdf_txt_path)
    texts_sheet2=texts_sheet2.split("\n")

    model =SentenceTransformer('bert-base-nli-mean-tokens')

    #sheet1 = pd.read_excel(input_excel_path)
    texts_sheet1 = sheet1[sheet1['Answer'] == 'Not Answered']['Question'].tolist()
    
    #data = open('/content/Venus-SOC2-Type2 AN.pdf', 'rb')
    #with open('/content/drive/MyDrive/pytesseract text files/Venus-SOC2-Type2 AN.txt', 'r') as file:
        #texts_sheet2 = file.readlines()
    #print(texts_sheet2)
    #print(texts_sheet1)
    embeddings_sheet1 = model.encode(texts_sheet1)
    embeddings_sheet2 = model.encode(texts_sheet2)
    print(embeddings_sheet1.shape)
    print(embeddings_sheet2.shape)
    matching_questions = []
    unmatched_questions = []

    for i, embedding_sheet1 in enumerate(embeddings_sheet1):
        max_similarity = 0.0
        max_similarity_index = -1
        for j, embedding_sheet2 in enumerate(embeddings_sheet2):
            similarity = cosine_similarity([embedding_sheet1], [embedding_sheet2])[0][0]
            if similarity >= 0.75 and similarity > max_similarity:
                max_similarity = similarity
                max_similarity_index = j

        if max_similarity_index != -1:
            matching_questions.append((texts_sheet1[i], texts_sheet2[max_similarity_index], max_similarity, 'Yes'))
            sheet1.loc[sheet1[sheet1['Question'] == texts_sheet1[i]].index, 'Answer'] = 'Yes'  # Update 'Answer' column to 'Yes'
        else:
            unmatched_questions.append(texts_sheet1[i])
            sheet1.loc[sheet1[sheet1['Question'] == texts_sheet1[i]].index, 'Answer'] = 'No'
        return sheet1 # .to_json()  # Update 'Answer' column to 'No'


# @app.route("/combined_compare_questionnaires", methods=["GET", "POST"])
# @cross_origin()
def combined_compare_questionnaires():
    print("in this function")
    data = request.files['SIGlight']
    data.save(os.path.join(f'{proot}/temp/', data.filename))
    file_path = os.path.join(f'{proot}/temp/', data.filename)
    # fileobj = data.read()

    base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions1 = pd.read_excel(base_file)

    # Set up base template
    base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions = pd.read_excel(base_file)
    base_questions = base_questions['Controls'][1:]
    print(base_questions)

    base_questions = base_questions.apply(lambda row: re.sub("^\d+\.\d+", "", row, count=1))

    # Set up file to compare
    # new_file = file_path
    try:
        new_questions = pd.read_excel(file_path)
    except:
        return "Error in reading Excel file"
    file = pd.ExcelFile(file_path)
    sheets = file.sheet_names

    if len(sheets) > 1:
        if "Assessment details" in sheets:
            new_questions = assessment_details(sheets, file_path,data.filename,base_questions1)
        elif "Cover Page" in sheets:
            new_questions = cover_page(sheets, file_path,data.filename,base_questions1)
    else:
        new_questions = sheets_zero(file_path,data.filename,base_questions1)


    dict = {'Question': [],
            'Answer': []}

    for row in base_questions:
        temp = process.extractOne(row, new_questions['Question'])
        if temp[1] > 85:
            dict['Question'].append(row)
            if type(new_questions['Answer'][temp[2]]) is not float:
                dict['Answer'].append(new_questions["Answer"][temp[2]])
            else:
                dict['Answer'].append("Not Answered")
        else:
            dict['Question'].append(row)
            dict['Answer'].append("No")
    print("function end")
    df = pd.DataFrame(dict)  
    dfup = update_excel_with_similarity(df)
    updated_df=perform_sentiment_analysis(dfup)
    response_dict = {"Question": updated_df['Question'].to_list(), "Answer": updated_df['Answer'].to_list()}
    #response_dict = {"Question": dfup['Question'].to_list(), "Answer": dfup['Answer'].to_list()}
    #dfup.to_excel("/content/finaloutput_withupdate.xlsx")
    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # updated_df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()

    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))
    print("RESPONSE DICT")
    print(len(response_dict['Question']))
    print(len(base_questions))
    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    print([i.cells[1].text])
                    print(updated_df.index[updated_df['Question'] == i.cells[1].text])
                    # if not updated_df.index[updated_df['Question'] == i.cells[1].text].empty:
                    count = updated_df.index[updated_df['Question'] == i.cells[1].text][0]
                    c += 1
                    i.cells[2].text = updated_df['Answer'][count]

    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]
        
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if 'Vendor Name' in run.text:
                run.text = run.text.replace(old_word, new_word)
                break

    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")              
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\")


    # updated_df.to_excel(f"{proot}/output/output.xlsx")
    #dfup.to_excel(f"{proot}/output/test2.xlsx")
    return jsonify(response_dict)# dfup.to_json() # json.dumps(dict, indent=4)


@app.route("/process_files", methods=["POST", "GET"])
@cross_origin()
def process_files():
    keys = request.files.keys()
    if "SOC2" in keys and "SIGlight" in keys:
        # return "both are here"
        return combined_compare_questionnaires()
    elif "SOC2" in keys:
        # return "SOC2 is here"
        return compare_pdfs()
    elif "SIGlight" in keys:
        # return "SIGlight is here"
        return compare_questionnaires()
    else:
        return "Invalid Request"

'''@app.route("/download_report", methods=["POST", "GET"])
@cross_origin()
def download_report():
    # with open(os.path.join('output', 'output.xlsx'), 'r') as f:
    #     resp = Response(f.read())

    # # set headers to tell encoding and to send as an attachment
    # resp.headers["Content-Encoding"] = 'gzip'
    # resp.headers["Content-Disposition"] = "attachment; filename=output.xlsx"
    # resp.headers["Content-type"] = "text/csv"
    #return send_file(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", 'D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.pdf'), as_attachment=True, download_name="report.pdf")
    # return resp
    return send_file(os.path.join('output', 'output.pdf'), as_attachment=True, download_name="report.pdf")'''


@app.route("/download_report", methods=["POST", "GET"])
@cross_origin()
def download_report():
    # Generate the report
   

    return send_file(os.path.join('output', 'output.pdf'), mimetype='application/pdf')  # Redirect to the route that handles the file download





# driver function
if __name__ == '__main__':
    app.run(host="0.0.0.0", debug=True)
