import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
#from fuzzywuzzy import fuzz
# from fuzzywuzzy import process
from flask import send_file
#from rapidfuzz import fuzz
#from rapidfuzz import process
import os
#import pytesseract
from pdf2image import convert_from_path
from flask import Flask, request, jsonify, send_from_directory, send_file, Response, url_for
from flask_cors import CORS, cross_origin

#from sklearn.metrics.pairwise import cosine_similarity
#from sentence_transformers import SentenceTransformer

from datetime import datetime
from textblob import TextBlob
import json
import re
import PyPDF2
from docx import Document
from config import *
import subprocess

def generate_pdf(doc_path, path):
    print("inside PDF func")
    
    print(doc_path)
    print(path)
    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path

app = Flask(__name__)
app.config['CORS_HEADERS'] = 'Content-Type'
CORS(app, expose_headers=["Content-Disposition"])
# remove the "/" in the end of path
proot =BASE_DIR

if not os.path.exists(f'{proot}/temp'):
    with open(f'{proot}/temp', "w"): pass

name_dict = {'galaxy Siglight-AN.xlsx': "Galaxy",
            'Jupiter Siglight-AN.xlsx': "Jupiter",
            'Mars- SIG Lite-AN.xlsx': "Mars",
            "Orit Siglight-AN.xlsx": "Orbit",
            "SIG Lite- Vendor Alpha-Anita.xlsx": "Alpha",
            "SIG Lite- Vendor Beta-Anita.xlsx": "Beta",
            "Vendor Sigma SIG Lite-Anita.xlsx": "Sigma",
            "Vendor Froce_SIG Lite _ AN.xlsx": "Force",
            "Solar-Siglight_AN.xlsx": "Solar"
            }


'''def merge_sheet_answer(column_10_data,column_11_data,column_12_data,column_13_data,new_questions,j):
    
    new_data = {'Question': [], 'Answer': [], 'Risk': []}
    if ('Sheet' not in new_questions.columns) and ('Serial number' not in new_questions.columns):
        
        #print(new_questions["Question"][4])
        for i in range(len(column_13_data)):
            answer = None
            #print("question length and row length")
            #print(len(column_12_data))
            #print(len(column_13_data))
            value_12=column_12_data[i]
            value_13 = column_13_data[i]
            
            new_data['Question'].append(value_13)
            match_index = new_questions.index[((new_questions.index)+1) == value_12].tolist()
            #print("match index")
            #print(match_index)
            if match_index:
                
                answer = new_questions.loc[match_index[0], 'Answer']
                new_data['Answer'].append(answer)
                
            else:
                new_data['Answer'].append('Not Available')
            

            answer_str = str(answer)  

            if answer_str.lower() in ['yes', 'yes ', 'Yes', 'Yes ', "YES"]:
                risk = 'Green'
            elif answer_str.lower() in ['no', 'no ', 'No', 'No ', "NO"]:
                risk = 'Red'
            else:
                risk = ''

            new_data['Risk'].append(risk)
        #print(new_questions1)
        new_questions1 = pd.DataFrame(new_data)
        print("final result")
        print(new_questions1)
        return new_questions1
    else:    
        
        for i in range(len(column_13_data)):
            value_10 = column_10_data[i]
            value_11 = column_11_data[i]
            value_12 = column_12_data[i]
            value_13 = column_13_data[i]
            if value_10 not in ['yes', 'Yes', 'no', 'No', 'yes ', 'Yes ', 'no ', 'No ']:
                new_data['Question'].append(value_13)
                new_data['Answer'].append('Not Available')
                new_data['Risk'].append('')
                continue
            if j==10:
                match = new_questions[(new_questions['Serial number'] == str(value_12))]
            elif j==6:
                match = new_questions[(new_questions['Sheet'] == value_11) & ((new_questions['Serial number']) == value_12-3)]   
            else:
                match = new_questions[(new_questions['Sheet'] == value_11) & (new_questions['Serial number'] == value_12)]
            if not match.empty:
                answer = str(match.iloc[0]['Answer'])
        
                if answer.lower() in ['yes', 'yes ', 'Yes', 'Yes ', 'YES']:
                    risk = 'Green'
           
                elif answer.lower() in ['no', 'no ', 'No', 'No ', 'NO']:
                    risk = 'Red'
                else:
                    risk = ''
                new_data['Question'].append(value_13)
                new_data['Answer'].append(answer)
                new_data['Risk'].append(risk)
        new_questions1 = pd.DataFrame(new_data)
        print("final result")
        print(new_questions1)
        return new_questions1'''


def assessment_details(sheets, file_path):
    sheets.remove("Assessment details")
    df = pd.concat(pd.read_excel(file_path, sheet_name=sheets))
    new_questions = pd.DataFrame({"Question": df['Question'].to_list(), "Answer": df['Actual answers'].to_list()})
    new_questions['Question'] = new_questions['Question'].apply(lambda row: re.sub("^\d+\.\d+|^[A-Z]{2}\.\d+\.", "", row, count=1))
    return new_questions

def cover_page_no_lite(sheets, file_path):
    print("in cover page no lite")
    remove_sheet = ['Cover Page', 'Business Information', 'Documentation', 'Lite', 'Glossary', 'Version History', 'Formula Notes', 'Full', 'Copyright', 'Terms of Use']
    for sheet in remove_sheet:
        sheets.remove(sheet)
    df = pd.concat(pd.read_excel(file_path, sheet_name=sheets, index_col=None, usecols="C,D" ))
    df = df.reset_index()
    new_questions = pd.DataFrame({"Question": df['Unnamed: 2'][2:].to_list(), "Answer": df['Unnamed: 3'][2:].to_list()})
    #if filename=="Orbit Siglight-AN.xlsx":
    
    #new_questions.to_excel("hello2.xlsx")
    return new_questions
    



def cover_page(sheets, file_path):
    df = pd.read_excel(file_path, sheet_name='Lite')
    for i, j in enumerate(df['Unnamed: 3']):
        if j == 'Response':
            if pd.isna(df['Unnamed: 3'][i+1]) and pd.isna(df['Unnamed: 3'][i+2]) and pd.isna(df['Unnamed: 3'][i+3]) and pd.isna(df['Unnamed: 3'][i+4]):
                return cover_page_no_lite(sheets, file_path)
            else:
                new_questions = pd.DataFrame()
                new_questions['Question'] = df['Unnamed: 2'][2:]
                new_questions['Answer'] = df['Unnamed: 3'][2:]
                #df['SIG Lite'] = df['SIG Lite'].str.replace('SL.', '', regex=True).str.strip()
                #new_questions['Serial number'] = df['SIG Lite'][2:]
    
    #if filename=="Jupiter SIglight-AN.xlsx":
    
    return new_questions


def sheets_zero(file_path):
    start_row = 0
    start_column = 0

    new_questions = pd.read_excel(file_path, header=None)

    for index, row in new_questions.iterrows():
        if 'Question' in str(row):
            start_row = index
            start_column = row[row.astype(str).str.contains('Question', case=False, na=False)].index.tolist()[0]
            # print(start_row)
            # print(start_column)
            break
        elif 'Questionnaire' in str(row):
            continue
    questions = new_questions.iloc[start_row+1:, start_column]
    questions = questions.str.replace(r"^\d+\.\d+|^[A-Z]{2}\.\d+\.", '', regex=True)
    #print(questions)

    for index, row in new_questions.iterrows():
        if 'Response' in str(row):
            start_row = index
            start_column = row[row.astype(str).str.contains('Response', case=False, na=False)].index.tolist()[0]
            break

    answer = new_questions.iloc[start_row+1:, start_column]

    # new_questions = new_questions[['Unnamed: 0', 'Unnamed: 2']][7:]
    new_questions = pd.DataFrame({"Question": questions, "Answer": answer})
    new_questions['Question'] = new_questions['Question'].apply(lambda row: re.sub("^\d+\.\d+", "", str(row), count=1))

    return new_questions


'''def perform_sentiment_analysis(data):
    def sentiment_analysis(answer):
        try:
            if answer.startswith("not available") or answer.startswith("Not Available") or answer.startswith("NOT AVAILABLE") or answer.startswith("not available ") or answer.startswith("Not Available ") or answer.startswith("NOT AVAILABLE "):
                sentiment_label = 'Not Available'
                polarity = None
                #subjectivity = None
            if answer.startswith("Yes") or answer.startswith("yes") or answer.startswith("YES") or answer.startswith("Yes ") or answer.startswith("yes ") or answer.startswith("YES "):
                sentiment_label = 'Yes'
                polarity = None
                #subjectivity = None
            
            elif answer.startswith("No") or answer.startswith("NO") or answer.startswith("no") or answer.startswith("No ") or answer.startswith("NO ") or answer.startswith("no "):
                sentiment_label = 'No'
                polarity = None
                #subjectivity = None
            elif answer.startswith("N/A") or answer.startswith("n/a") or answer.startswith("Not Answered") or answer.startswith("NA") or answer.startswith("NOT ANSWERED") or answer.startswith("not answered") or answer.startswith("Not answered") or answer.startswith("na") or answer.startswith("N A"):
                sentiment_label='Not Answered'
                polarity=None
            else:
                blob = TextBlob(answer)
                polarity = blob.sentiment.polarity
                #subjectivity = blob.sentiment.subjectivity

                if polarity >= 0:
                    sentiment_label = 'Yes'
                elif polarity < 0:
                    sentiment_label = 'No'
        except:
            sentiment_label = 'Not Available'
            polarity = None

        return sentiment_label

    data[['Answer']] = data['Answer'].apply(sentiment_analysis).apply(pd.Series)

    return data'''
def perform_sentiment_analysis(data):
    # Define a dictionary to map starting keywords to sentiment labels
    sentiment_mapping = {
        'not available': 'Not Available',
        'yes': 'Yes',
        'no': 'No',
        'n/a': 'Not Answered',
        'not answered': 'Not Answered',
        'na': 'Not Answered',
        'n a': 'Not Answered'
    }

    # Function to get sentiment label for each answer
    def get_sentiment_label(answer):
        answer_lower = answer.lower().strip()
        #print("answerlower:", answer_lower)

        for keyword, sentiment_label in sentiment_mapping.items():
            keyword_pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
            if re.search(keyword_pattern, answer_lower):
                #print("Match found!")
                return sentiment_label

        # Perform sentiment analysis using TextBlob if no keyword match found
        blob = TextBlob(answer)
        polarity = blob.sentiment.polarity
        return 'Yes' if polarity >= 0 else 'No'

    # Update the 'Answer' column with the sentiment labels
    data['Answer'] = data['Answer'].apply(get_sentiment_label)
    
    return data

stopwords= ['i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', "you're", "you've",\
            "you'll", "you'd", 'your', 'yours', 'yourself', 'yourselves', 'he', 'him', 'his', 'himself', \
            'she', "she's", 'her', 'hers', 'herself', 'it', "it's", 'its', 'itself', 'they', 'them', 'their',\
            'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', "that'll", 'these', 'those', \
            'am', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', \
            'did', 'doing', 'a', 'an', 'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', \
            'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into', 'through', 'during', 'before', 'after',\
            'above', 'below', 'to', 'from', 'up', 'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further',\
            'then', 'once', 'here', 'there', 'when', 'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more',\
            'most', 'other', 'some', 'such', 'only', 'own', 'same', 'so', 'than', 'too', 'very', \
            's', 't', 'can', 'will', 'just', 'don', "don't", 'should', "should've", 'now', 'd', 'll', 'm', 'o', 're', \
            've', 'y', 'ain', 'aren', "aren't", 'couldn', "couldn't", 'didn', "didn't", 'doesn', "doesn't", 'hadn',\
            "hadn't", 'hasn', "hasn't", 'haven', "haven't", 'isn', "isn't", 'ma', 'mightn', "mightn't", 'mustn',\
            "mustn't", 'needn', "needn't", 'shan', "shan't", 'shouldn', "shouldn't", 'wasn', "wasn't", 'weren', "weren't", \
            'won', "won't", 'wouldn', "wouldn't"]

'''def jaccard_similarity(str1, str2):
    #if isinstance(str2, float):
        #str2 = str(str2)
    str1=str(str1)
    str2=str(str2)
    set1 = set(str1.lower().split())
    set2 = set(str2.lower().split())
    intersection = set1.intersection(set2)
    union = set1.union(set2)
    return len(intersection) / len(union)'''
def clean_text(text):
    # Remove digits and punctuation, and convert to lowercase
    text = re.sub(r'\d+', '', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text.lower()

def jaccard_similarity(str1, str2):
    str1 = clean_text(str1)
    str2 = clean_text(str2)
    
    words_set1 = set(str1.split()) - set(stopwords)
    words_set2 = set(str2.split()) - set(stopwords)
    
    intersection = words_set1.intersection(words_set2)
    union = words_set1.union(words_set2)
    
    return len(intersection) / len(union)
#import Levenshtein

# Update the similarity function to use Levenshtein Distance
#def levenshtein_similarity(str1, str2):
    #return 1 - Levenshtein.distance(str1, str2) / max(len(str1), len(str2))



@app.route("/compare_questionnaires", methods=["GET", "POST"])
@cross_origin()
def compare_questionnaires():
    print("in this function")
    data = request.files['SIGlight']
    data.save(os.path.join(f'{proot}/temp/', data.filename))
    file_path = os.path.join(f'{proot}/temp/', data.filename)
    # fileobj = data.read()

    # Set up base template
    base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions = pd.read_excel(base_file)
    base_questions= base_questions.iloc[1:, 2].reset_index(drop=True)
    base_questions = base_questions.apply(lambda row: re.sub("^\d+\.\d+", "", row, count=1))

    # Set up file to compare
    # new_file = file_path
    try:
        new_questions = pd.read_excel(file_path)
    except:
        return "Error in reading Excel file"
    file = pd.ExcelFile(file_path)
    sheets = file.sheet_names

    if len(sheets) > 1:
        if "Assessment details" in sheets:
            new_questions = assessment_details(sheets, file_path)
        elif "Cover Page" in sheets:
            new_questions = cover_page(sheets, file_path)
    else:
        new_questions = sheets_zero(file_path)

    #new_questions['Question'] = new_questions['Question'].apply(lambda row: ' '.join([word for word in str(row).split() if word.lower() not in stopwords]))  #new line

    dict = {'Question': [],
            'Answer': []}
    

    '''matched_new_indices = set()

   

    base_questions_df = pd.DataFrame({'Question': base_questions})  # Convert Series to DataFrame

    for _, base_row in base_questions_df.iterrows():
        best_similarity = 0.0
        second_best_similarity = 0.0
        best_new_index = None
        second_best_new_index = None

        for new_index, new_row in new_questions.iterrows():
            # Skip new questions that have already been matched
            if new_index in matched_new_indices:
                continue

            similarity = jaccard_similarity(base_row['Question'], new_row['Question'])

            if similarity > best_similarity:
                # Update best and second best similarities and their respective indices
                second_best_similarity = best_similarity
                second_best_new_index = best_new_index

                best_similarity = similarity
                best_new_index = new_index
            elif similarity > second_best_similarity:
                # Update only second best similarity and its index
                second_best_similarity = similarity
                second_best_new_index = new_index

        if best_similarity > 0.1 and best_new_index is not None:
            matched_new_indices.add(best_new_index)

            print("base questions", base_row['Question'])
            dict['Question'].append(base_row['Question'])
            answer_value = new_questions.loc[best_new_index, 'Answer']
            if not pd.isna(answer_value) and str(answer_value).strip() != '':
                dict['Answer'].append(answer_value)
            else:
                dict['Answer'].append("Not Answered")
            print("matching question:", new_questions.loc[best_new_index, 'Question'])
        elif second_best_similarity > 0.1 and second_best_new_index is not None:
            matched_new_indices.add(second_best_new_index)

            print("base questions", base_row['Question'])
            dict['Question'].append(base_row['Question'])
            answer_value = new_questions.loc[second_best_new_index, 'Answer']
            if not pd.isna(answer_value) and str(answer_value).strip() != '':
                dict['Answer'].append(answer_value)
            else:
                dict['Answer'].append("Not Answered")
            print("matching question:", new_questions.loc[second_best_new_index, 'Question'])
        else:
            dict['Question'].append(base_row['Question'])
            dict['Answer'].append("Not Available")'''
    



    for row in base_questions:
        
        max_similarity = 0.1
        best_new_index = None

        for new_index, new_row in new_questions.iterrows():
            similarity = jaccard_similarity(row, new_row['Question'])
            if similarity > max_similarity:
                max_similarity = similarity
                best_new_index = new_index

        if max_similarity > 0.1 and best_new_index is not None:
            dict['Question'].append(row)
            answer_value = new_questions.loc[best_new_index, 'Answer']
            if not pd.isna(answer_value) and str(answer_value).strip() != '':
                dict['Answer'].append(new_questions["Answer"][best_new_index])
            else:
                dict['Answer'].append("Not Answered")
        else:
            dict['Question'].append(row)
            dict['Answer'].append("Not Available")

    print("function end")
    #print(len(dict['Question']))
    #print(len(dict['Answer']))
    '''matched_new_indices = set()

    for row in base_questions:
        max_similarity = 0.1
        best_new_index = None

        for new_index, new_row in new_questions.iterrows():
            # Skip new questions that have already been matched
            if new_index in matched_new_indices:
                continue

            similarity = levenshtein_similarity(row, new_row['Question'])

            if similarity > max_similarity:
                max_similarity = similarity
                best_new_index = new_index

        if max_similarity > 0.1 and best_new_index is not None:
            matched_new_indices.add(best_new_index)

            print("base questions", row)
            dict['Question'].append(row)
            answer_value = new_questions.loc[best_new_index, 'Answer']
            if not pd.isna(answer_value) and str(answer_value).strip() != '':
                dict['Answer'].append(answer_value)
            else:
                dict['Answer'].append("Not Answered")
            print("matching question:", new_questions.loc[best_new_index, 'Question'])
        else:
            dict['Question'].append(row)
            dict['Answer'].append("Not Available")'''




    df = pd.DataFrame(dict) 
    #print("**********************BEFORE SENTIMENT",df)
    #print(type(df.iloc[0]))
    updated_df = perform_sentiment_analysis(df)
    #print("44444444444AFTER SENTIMENT",updated_df)



    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # updated_df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()
    # # updated_df.to_excel(f"{proot}/output/output.xlsx")
    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))

    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    #print([i.cells[1].text])
                    #print(updated_df.index[updated_df['Question'] == i.cells[1].text])
                    # if not updated_df.index[updated_df['Question'] == i.cells[1].text].empty:
                    count = updated_df.index[updated_df['Question'] == i.cells[1].text][0]
                    c += 1
                    i.cells[2].text = updated_df['Answer'][count]

    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "Vendor Name" in run.text:
                run.text = run.text.replace(old_word, new_word)
                break
    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\")

    return json.dumps({'Question': updated_df['Question'].to_list(),
                       'Answer': updated_df['Answer'].to_list()}, indent=4)

def convert_pdf_to_text(pdf_file):
    # pdf_file = open("soc_2/Venus-SOC2-Type2 AN.pdf", "rb")
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    num_pages = len(pdf_reader.pages)
    txt_file = open("temp_pdf_text.txt", "w",encoding='utf-8')
    for i in range(num_pages):
        page_text = pdf_reader.pages[i].extract_text()
        txt_file.write(page_text)
        # print(page_text)
    # pdf_file.close()
    txt_file.close()
    return "temp_pdf_text.txt"



def clean_pdf(pdf_txt_path):
    with open(pdf_txt_path, 'r', encoding='utf-8') as file:
        texts = file.readlines()

    cleaned_lines = []
    for text in texts:
        text = text.strip()

        if not text:
            continue

        text = re.sub(r'\b\w\b', '', text)
        text = re.sub(r"\\", " ", text)
        text = re.sub(r"\/", " ", text)
        text = re.sub(r"[\n\t\-]", " ", text)
        text = re.sub(r"won't", "will not", text)
        text = re.sub(r"won’t", "will not", text)
        text = re.sub(r"can\'t", "can not", text)
        text = re.sub(r"can\’t", "can not", text)
        text = re.sub(r"n\'t", " not", text)
        text = re.sub(r"n\’t", " not", text)
        text = re.sub(r"\'re", " are", text)
        text = re.sub(r"\’re", " are", text)
        text = re.sub(r"\'s", " is", text)
        text = re.sub(r"\’s", " is", text)
        text = re.sub(r"\'d", " would", text)
        text = re.sub(r"\’d", " would", text)
        text = re.sub(r"\'ll", " will", text)
        text = re.sub(r"\’ll", " will", text)
        text = re.sub(r"\'t", " not", text)
        text = re.sub(r"\’t", " not", text)
        text = re.sub(r"\'ve", " have", text)
        text = re.sub(r"\’ve", " have", text)
        text = re.sub(r"\’m", " am", text)
        text = re.sub(r"\'m", " am", text)
        text = re.sub(r'[^A-Za-z_]', " ", text)
        text = re.sub(r"\s+", " ", text)
        text = ' '.join(e for e in text.split() if e.lower() not in stopwords)
        cleaned_lines.append(text.lower())

    cleaned_text = "\n".join(cleaned_lines)
    if os.path.exists(pdf_txt_path):
        os.remove(pdf_txt_path)
    return cleaned_text

def split_into_lines(text, words_per_line=12):
    words = text.split()
    lines = []

    current_line = []
    for word in words:
        if len(current_line) + len(word.split()) <= words_per_line:
            current_line.extend(word.split())
        else:
            lines.append(current_line)
            current_line = word.split()

    # Append any remaining words to the last line
    if current_line:
        lines.append(current_line)

    lines_with_exactly_words = []

    for line in lines:
        while len(line) < words_per_line:
            line.append('')  # Add empty strings to pad the line to exactly 25 words
        lines_with_exactly_words.append(line)

    return "\n".join([" ".join(line) for line in lines_with_exactly_words])

@app.route("/compare_pdfs", methods=['POST', 'GET'])
@cross_origin()
def compare_pdfs():
    # Load the BERT model
    data = request.files['SOC2']
    sheet1 = pd.read_excel(f'{proot}/base_files/template(Compare_pdf).xlsx')
    # print(data.read())
    file_name = data.filename
    pdf_path = os.path.join(f"{proot}/temp/", file_name)
    #data.save(pdf_path)   #change
    #output_text_file = os.path.join("output/", str(datetime.now()))  #ch
    #output_text_file = os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\", str(datetime.now()))
    pdf_txt_path = convert_pdf_to_text(data)
    texts_sheet2 = clean_pdf(pdf_txt_path)
    texts_sheet2 = split_into_lines(texts_sheet2, words_per_line=12)
    texts_sheet2 = texts_sheet2.split('\n')
    texts_sheet1 = sheet1['questions'].tolist()
    texts_sheet1 = [text.lower() for text in texts_sheet1]

    matching_questions = []
    unmatched_questions = []

    for i, text_sheet1 in enumerate(texts_sheet1):
        max_similarity = 0.0
        max_similarity_index = -1
        for j, text_sheet2 in enumerate(texts_sheet2):

            # Convert texts to sets of words
            words_set1 = set(text_sheet1.split())
            words_set2 = set(text_sheet2.split())

            # Calculate Jaccard similarity between the sets
            jaccard_similarity = len(words_set1.intersection(words_set2)) / len(words_set1.union(words_set2))
            #print(jaccard_similarity)
            if jaccard_similarity >= 0.1 and jaccard_similarity > max_similarity:
                max_similarity = jaccard_similarity
                max_similarity_index = j

        if max_similarity_index != -1:
            matching_questions.append((sheet1.iloc[i, 0], texts_sheet2[max_similarity_index], max_similarity, 'Yes'))
        else:
            unmatched_questions.append(sheet1.iloc[i, 0])

    matching_df = pd.DataFrame(matching_questions, columns=['Question 1', 'Question 2', 'Similarity', 'Matching'])
    updated_df = pd.merge(sheet1, matching_df, left_on='questions', right_on='Question 1', how='left')

    updated_df['Question 2'].fillna('', inplace=True)
    updated_df['Similarity'].fillna('', inplace=True)
    updated_df['Matching'].fillna('No', inplace=True)

    updated_df = updated_df.sort_values(by='questions', key=lambda x: x.map({q: i for i, q in enumerate(sheet1['questions'])}))
    updated_df.drop(columns=['Question 1'], inplace=True)

    updated_df.rename(columns={'questions': 'Questionnaires', 'Question 2': 'POC 2 DATA'}, inplace=True)

    updated_df1 = {'Question': updated_df['Questionnaires'].to_list(),
                  'Answer': updated_df['Matching'].to_list()}

    df = pd.DataFrame(updated_df1)
    # excel_file = file_name.split(".")[0] + ".xlsx"
    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()
    #print(df)
    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))

    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    #print([i.cells[1].text])
                    #print(df.index[df['Question'] == i.cells[1].text])
                    i_cells_text = i.cells[1].text.strip()
                        # Apply strip() to the 'Question' column
                    df['Question'] = df['Question'].str.strip()
                    count = df.index[df['Question'] == i_cells_text][0]
                    c += 1
                    i.cells[2].text = df['Answer'][count]
    
    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "Vendor Name" in run.text:
                run.text = run.text.replace(old_word, new_word)
                break


    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output")


    # matching_df.to_excel(os.path.join("output/", "output.xlsx"), index=False)
    # updated_df.to_excel(os.path.join("temp/", "output.xlsx"), index=False)
    return json.dumps(updated_df1, indent=4)
    # print("Unmatched Questions:")
    # for question in unmatched_questions:
    #     print(question)


def update_excel_with_similarity(sheet1):
    if 'Not Answered' not in sheet1['Answer'].values:
        #print("in not answered")
        return sheet1
    data = request.files['SOC2']
    data2 = request.files['SIGlight']
    file_name = data.filename
    pdf_path = os.path.join(f"{proot}/temp/", file_name)
    data.save(pdf_path)
    output_text_file = os.path.join("output/", str(datetime.now()))
    #output_text_file = os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", str(datetime.now()))
    pdf_txt_path = convert_pdf_to_text(data)
    texts_sheet2 = clean_pdf(pdf_txt_path)
    texts_sheet2 = split_into_lines(texts_sheet2, words_per_line=12)
    texts_sheet2 = texts_sheet2.split('\n')
    filtered_indices = sheet1.index[sheet1['Answer'] == 'Not Answered']
    filtered_sheet1 = sheet1.loc[filtered_indices].copy()
    #print("naaaa", filtered_sheet1)
    texts_sheet1 = filtered_sheet1['Question'].tolist()  # Use the "Question" column for comparison
    texts_sheet1 = [text.lower() for text in texts_sheet1]
    # model = SentenceTransformer('bert-base-nli-mean-tokens')
    for i, text_sheet1 in enumerate(texts_sheet1):
        max_similarity = 0.0
        max_similarity_index = -1
        for j, text_sheet2 in enumerate(texts_sheet2):

            # Convert texts to sets of words
            words_set1 = set(text_sheet1.split())
            words_set2 = set(text_sheet2.split())

            # Calculate Jaccard similarity between the sets
            jaccard_similarity = len(words_set1.intersection(words_set2)) / len(words_set1.union(words_set2))

            if jaccard_similarity > max_similarity:
                max_similarity = jaccard_similarity
                max_similarity_index = j

        if max_similarity_index != -1 and max_similarity > 0.1:
            sheet1.at[filtered_indices[i], 'Answer'] = 'Yes'
        else:
            sheet1.at[filtered_indices[i], 'Answer'] = 'No'

    #print("sheet1 final", sheet1)
    return sheet1 # .to_json()  # Update 'Answer' column to 'No'


# @app.route("/combined_compare_questionnaires", methods=["GET", "POST"])
# @cross_origin()
def combined_compare_questionnaires():
    print("in this function")
    data = request.files['SIGlight']
    data.save(os.path.join(f'{proot}/temp/', data.filename))
    file_path = os.path.join(f'{proot}/temp/', data.filename)
    # fileobj = data.read()

    '''base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions1 = pd.read_excel(base_file)'''

    # Set up base template
    base_file = f"{proot}/base_files/base_template.xlsx"
    base_questions = pd.read_excel(base_file)
    base_questions= base_questions.iloc[1:, 2].reset_index(drop=True)
    base_questions = base_questions.apply(lambda row: re.sub("^\d+\.\d+", "", row, count=1))

    # Set up file to compare
    # new_file = file_path
    try:
        new_questions = pd.read_excel(file_path)
    except:
        return "Error in reading Excel file"
    file = pd.ExcelFile(file_path)
    sheets = file.sheet_names

    if len(sheets) > 1:
        if "Assessment details" in sheets:
            new_questions = assessment_details(sheets, file_path)
        elif "Cover Page" in sheets:
            new_questions = cover_page(sheets, file_path)
    else:
        new_questions = sheets_zero(file_path)

    #new_questions['Question'] = new_questions['Question'].apply(lambda row: ' '.join([word for word in str(row).split() if word.lower() not in stopwords]))  #new line

    dict = {'Question': [],
            'Answer': []}

    for row in base_questions:
        
        max_similarity = 0.1
        best_new_index = None

        for new_index, new_row in new_questions.iterrows():
            similarity = jaccard_similarity(row, new_row['Question'])
            if similarity > max_similarity:
                max_similarity = similarity
                best_new_index = new_index

        if max_similarity > 0.1 and best_new_index is not None:
            dict['Question'].append(row)
            answer_value = new_questions.loc[best_new_index, 'Answer']
            if not pd.isna(answer_value) and str(answer_value).strip() != '':
                dict['Answer'].append(new_questions["Answer"][best_new_index])
            else:
                dict['Answer'].append("Not Answered")
        else:
            dict['Question'].append(row)
            dict['Answer'].append("Not Available")
    #print("function end")
    df = pd.DataFrame(dict)  
    dfup=perform_sentiment_analysis(df)
    #print("dffdfdf0",df)
    updated_df = update_excel_with_similarity(dfup)
    #print("dfup",dfup)
    

    #response_dict = {"Question": dfup['Question'].to_list(), "Answer": dfup['Answer'].to_list()}
    #dfup.to_excel("/content/finaloutput_withupdate.xlsx")
    # # Create a Pandas Excel writer using XlsxWriter as the engine.
    # writer = pd.ExcelWriter(os.path.join("output/", "output.xlsx"), engine='xlsxwriter')

    # # Convert the dataframe to an XlsxWriter Excel object.
    # updated_df.to_excel(writer, sheet_name='Sheet1', startrow=3)

    # # Get the xlsxwriter workbook and worksheet objects.
    # workbook  = writer.book
    # worksheet = writer.sheets['Sheet1']

    # # Insert an image.
    # worksheet.insert_image('A1', 'beaconer1.png')

    # # Close the Pandas Excel writer and output the Excel file.
    # writer.close()

    document = Document(os.path.join(f"{proot}/base_files", 'Report Template for Demo (copy 2).docx'))
    #print("RESPONSE DICT")
    #print(len(response_dict['Question']))
    #print(len(base_questions))
    tl = document.tables
    headings = []
    c = 0
    for a in tl:
        if a == tl[-1]:
            break
        else:
            for idx, i in enumerate(a.rows):
                if idx <= 0:
                    continue
                else:
                    #print([i.cells[1].text])
                    #print(updated_df.index[updated_df['Question'] == i.cells[1].text])
                    # if not updated_df.index[updated_df['Question'] == i.cells[1].text].empty:
                    count = updated_df.index[updated_df['Question'] == i.cells[1].text][0]
                    c += 1
                    i.cells[2].text = updated_df['Answer'][count]

    old_word = "Vendor Name"
    try:
        new_word = name_dict[data.filename]
    except:
        new_word = data.filename.split(".")[0]
        
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if 'Vendor Name' in run.text:
                run.text = run.text.replace(old_word, new_word)
                break

    document.save(os.path.join("output", "output.docx"))
    generate_pdf(os.path.join("output", "output.docx"), "output/")              
    #document.save(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.docx"))
    #generate_pdf(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", "output.docx"), "D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output\\")


    # updated_df.to_excel(f"{proot}/output/output.xlsx")
    #dfup.to_excel(f"{proot}/output/test2.xlsx")
    #return jsonify(response_dict)# dfup.to_json() # 
    return json.dumps({"Question": updated_df['Question'].to_list(), "Answer": updated_df['Answer'].to_list()}, indent=4)



@app.route("/process_files", methods=["POST", "GET"])
@cross_origin()
def process_files():
    keys = request.files.keys()
    if "SOC2" in keys and "SIGlight" in keys:
        # return "both are here"
        return combined_compare_questionnaires()
    elif "SOC2" in keys:
        # return "SOC2 is here"
        return compare_pdfs()
    elif "SIGlight" in keys:
        # return "SIGlight is here"
        return compare_questionnaires()
    else:
        return "Invalid Request"

'''@app.route("/download_report", methods=["POST", "GET"])
@cross_origin()
def download_report():
    # with open(os.path.join('output', 'output.xlsx'), 'r') as f:
    #     resp = Response(f.read())

    # # set headers to tell encoding and to send as an attachment
    # resp.headers["Content-Encoding"] = 'gzip'
    # resp.headers["Content-Disposition"] = "attachment; filename=output.xlsx"
    # resp.headers["Content-type"] = "text/csv"
    #return send_file(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", 'D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.pdf'), as_attachment=True, download_name="report.pdf")
    # return resp
    return send_file(os.path.join('output', 'output.pdf'), as_attachment=True, download_name="report.pdf")'''
@app.route("/download_report", methods=["POST", "GET"])
@cross_origin()
def download_report():
    # with open(os.path.join('output', 'output.xlsx'), 'r') as f:
    #     resp = Response(f.read())

    # # set headers to tell encoding and to send as an attachment
    # resp.headers["Content-Encoding"] = 'gzip'
    # resp.headers["Content-Disposition"] = "attachment; filename=output.xlsx"
    # resp.headers["Content-type"] = "text/csv"
    #return send_file(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", 'D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output.pdf'), as_attachment=True, download_name="report.pdf")
    # return resp
    #return send_file(os.path.join('output', 'output.pdf'), as_attachment=True, download_name="report.pdf")
    #return send_file(os.path.join("D:\\JUPPITER AI LABS(anuragraiofficial321)\\beaconers\\output", 'output.pdf'), mimetype='application/pdf')
    return send_file(os.path.join("output",'output.pdf'), mimetype='application/pdf')
# driver function
if __name__ == '__main__':
    app.run(host="0.0.0.0", debug=True)
